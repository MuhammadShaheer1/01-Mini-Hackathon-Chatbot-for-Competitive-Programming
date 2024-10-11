from django.shortcuts import render,redirect
from django.http import JsonResponse
from django.http import HttpResponse


from django.contrib import auth
from django.contrib.auth.models import User
import openai
import markdown
from .models import Chat
import os
from openai import OpenAI
import re
from django.utils import timezone
from bs4 import BeautifulSoup
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import io

from pathlib import Path
from langchain_community.document_loaders import CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings


# doc_path = ["database.csv", "Leetcode_Questions.csv"]
# docs = []
# for doc_file in doc_path:
#   file_path = Path(doc_file)
#   print(doc_file)
#   if not file_path.exists():
#       print(f"File {doc_file} does not exist. Check the path.")
#       continue
#   try:
#     if doc_file.endswith(".csv"):
#       loader = CSVLoader(doc_file)
#     elif doc_file.endswith(".pdf"):
#       loader = PyPDFLoader(doc_file)
#     elif doc_file.endswith(".docx"):
#       loader = Docx2txtLoader(doc_file)
#     elif doc_file.endswith(".txt") or doc_file.endswith(".md"):
#       loader = TextLoader(doc_file)
#     else:
#       print(f"DocumentType {doc_file.type} not supported.")
#       continue
#     docs.extend(loader.load())
#   except Exception as e:
#     print(f"Error loading document {doc_file}: {e}")


# text_splitter = RecursiveCharacterTextSplitter(chunk_size = 5000,
#                                                chunk_overlap = 1000)

# documents_chunks = text_splitter.split_documents(docs)

#openai.api_key = os.environ['sk-gNp96fNbpb-_XUMldsABqJHLvHGwmII6WMet-6BBkTT3BlbkFJgF212KfTK8slEz9auOqS_xfG5-62u1GbtaORgjQOAA']
#openai.api_key = os.getenv("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")

#penai_api_key = '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' # Replace YOUR_API_KEY with your openai apikey 
#openai.api_key = openai_api_key 



client = OpenAI(
    api_key="e0bf8b477b694cb5b7d74f08db4770ea",
    base_url="https://api.aimlapi.com/v1",
)

chat_history = []
# def ask_openai(message):
#     while True:
#         user_input = input("You: ")

#     # Retrieve the most relevant document from the vector database
#         # result = vector_db.similarity_search(user_input, k=1)
#         # retrieved_content = result[0].page_content if result else "No relevant document found."

#     # Append the user input and retrieved document to the chat history
#         chat_history.append({"role": "user", "content": user_input})
#     # chat_history.append({"role": "assistant", "content": retrieved_content})
#         messages = chat_history.copy()
#     # Sending the conversation history to OpenAI API
#         response = client.chat.completions.create(
#             model="o1-preview",
#             messages=messages
#             ,max_tokens = 50000
#         )

#     # Get and print the assistant's reply
#         assistant_reply = response.choices[0].message.content
#     #print(f"Assistant: {assistant_reply}")

#     # Append assistant's reply to the chat history for context in the next loop
#         chat_history.append({"role": "assistant", "content": assistant_reply})

#         return assistant_reply


client = OpenAI(
    api_key="e0bf8b477b694cb5b7d74f08db4770ea",
    base_url="https://api.aimlapi.com",
)
def ask_openai(message):
    response = client.chat.completions.create(
        model="o1-mini",  # Use the model you're using (in this case, 'o1-mini')
        messages=[
            {"role": "user", "content": message},  # Take user input as the message
        ],
        max_tokens=20000,  # Set max tokens if needed
    )
    answer = response.choices[0].message.content
    
    return answer

#openai.api_key = os.getenv("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")

def chatbot(request):
    chats = Chat.objects.filter(user=request.user)


    if request.method == 'POST':
        #export Function if not work remove these two lines
        if request.POST.get('download') == 'true':

            return export_chat(request)


        message = request.POST.get('message')
        response = ask_openai(message)
        #formatted_response= "{response}"
        #Format Response Here

        response = markdown.markdown(response)
        print(response)
        chat = Chat(user=request.user, message=message, response=response, created_at=timezone.now)
        chat.save()
        #print(formatted_response)
        return JsonResponse({'message': message, 'response': response})
    
    return render(request, 'chatbot.html', {'chats': chats})


def login(request):
    if request.method=='POST':
        username = request.POST['username']
        password = request.POST['password']
        user = auth.authenticate(request, username=username, password=password)
        if user is not None:
            auth.login(request, user)
            return redirect('chatbot')
        else:
            error_message = 'Invalid username or password'
            return render(request, 'login.html', {'error_message': error_message})
    else:
        return render(request, 'login.html')

def register(request):
    if request.method == 'POST':
        username = request.POST['username']
        email = request.POST['email']
        password1 = request.POST['password1']
        password2 = request.POST['password2']

        if password1==password2:
            try:
                user = User.objects.create_user(username, email, password1)
                user.save()
                auth.login(request, user)
                return redirect('login')
            except:
                error_message = 'Error creating account'
            return render(request, 'register.html', {'error_message': error_message})
        else:
            error_message = "Password don't match" 
            return render(request, 'register.html', {'error_message': error_message})
    return render(request, 'register.html')

def logout(request):
    auth.logout(request)
    return redirect('register')


# Export Chat

'''
def export_chat(request):
    if request.user.is_authenticated:
        chats = Chat.objects.filter(user=request.user)
        
        # Prepare content for the .txt file
        content = ""
        for chat in chats:
            content += f"User: {chat.message}\nAI: {chat.response}\n\n"
        
        # Create an HttpResponse with a .txt file attachment
        response = HttpResponse(content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename=chat_history_{request.user.username}.txt'
        return response
    
    return redirect('chatbot')

'''

def export_chat(request):
    if request.user.is_authenticated:
        chats = Chat.objects.filter(user=request.user)
        
        # Create a new Word Document
        document = Document()
        
        # Add a title to the document
        document.add_heading('Chat History', level=1)
        
        for chat in chats:
            # Strip HTML tags from the message and response
            user_message = strip_html_tags(chat.message)
            ai_response = strip_html_tags(chat.response)
            
            # Add a heading for the user's message
            document.add_heading('User:', level=2)
            document.add_paragraph(user_message)
            
            # Add a heading for the AI's response
            document.add_heading('AI:', level=2)
            document.add_paragraph(ai_response)
            
            # Add line breaks between chats
            document.add_paragraph()
        
        # Save the document to an in-memory file
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        # Create an HttpResponse with a .docx file attachment
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=chat_history_{request.user.username}.docx'
        return response
    
    return redirect('chatbot')

def strip_html_tags(text):
    # Remove HTML tags using regular expressions
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)
